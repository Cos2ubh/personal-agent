"""
Document text extraction.

Pulls plain text out of PDFs, DOCX files, and text-family formats. Respects
the read sandbox and file-size cap from filesystem.py — we never touch a
document that read_file wouldn't touch.

The extracted text is what the LLM sees, so anything that improves signal
here (stripping headers/footers, joining hyphenated line breaks, ordering
columns) pays back on every downstream query. First pass just gets the raw
text out; refinements come later if quality is insufficient.
"""

from pathlib import Path

from tools.filesystem import _guard, MAX_READ_BYTES

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx",
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".csv", ".json", ".xml", ".yaml", ".yml",
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h",
}


def extract_text(path_str: str) -> str:
    """
    Extract plain text from a document. Returns text or an "Error: ..." string.
    Never raises — all failure modes come back as strings for LLM reasoning.
    """
    path = Path(path_str)

    # Sandbox check — same rules as read_file
    from tools.filesystem import PermissionDenied
    try:
        _guard(path, mode="read")
    except PermissionDenied as e:
        return f"PermissionDenied: {e}"

    if not path.exists():
        return f"Error: file not found — {path}"
    if not path.is_file():
        return f"Error: not a file — {path}"

    try:
        size = path.stat().st_size
    except OSError as e:
        return f"Error: cannot stat file — {e}"

    if size > MAX_READ_BYTES:
        return (
            f"Error: file too large — {size:,} bytes "
            f"(limit is {MAX_READ_BYTES:,} bytes / 10 MB)."
        )

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in SUPPORTED_EXTENSIONS:
        # Text-family formats: read directly
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            return f"Error reading text file: {e}"

    return (
        f"Error: unsupported file type '{ext}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF. Handles multi-page docs, falls back gracefully."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "Error: pypdf not installed. Run: pip install pypdf"

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        return f"Error opening PDF: {type(e).__name__}: {e}"

    if reader.is_encrypted:
        # Try empty password (many PDFs are technically encrypted but with no password)
        try:
            reader.decrypt("")
        except Exception:
            return "Error: PDF is password-protected."

    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            pages.append(f"[page {i + 1}: extraction failed — {e}]")
            continue
        if text.strip():
            pages.append(text.strip())

    if not pages:
        return "(no extractable text — may be a scanned image PDF; OCR not yet available)"

    return "\n\n---\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract text from a .docx file. Includes tables."""
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    try:
        doc = Document(str(path))
    except Exception as e:
        return f"Error opening DOCX: {type(e).__name__}: {e}"

    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        return "(no extractable text in this docx)"

    return "\n".join(parts)
